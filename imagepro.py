import cv2
import numpy as np
import pytesseract
from PIL import Image
import os
import hashlib
import docx

def generate_hashv(path):
    # Convert text to hash
    # img = np.array(img1)
    path1= './uploads-v/'+path



# Set Tesseract path
    os.environ['TESSDATA_PREFIX'] = 'C:/Users/Mahi Singhal/AppData/Local/Programs/Tesseract-OCR/tessdata'
    pytesseract.pytesseract.tesseract_cmd = r'C:/Users/Mahi Singhal/AppData/Local/Programs/Tesseract-OCR/tesseract.exe'

# Load image
    img = cv2.imread(path1)

# Convert image to grayscale
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

# Equalize histogram to enhance contrast
    equalized = cv2.equalizeHist(gray)

# Resize image
    new_size = (int(img.shape[1]*0.35), int(img.shape[0]*0.25))
    resized = cv2.resize(img, new_size) 

# Adjust brightness and contrast
    alpha = 1.25  # scaling factor for brightness adjustment
    beta = 30  # scalar added to the result
    brightened = cv2.convertScaleAbs(resized, alpha=alpha, beta=beta)

# Detect contours and find object with the biggest bounding box
    edges = cv2.Canny(gray, 50, 200)
    contours, hierarchy = cv2.findContours(edges.copy(), cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_NONE)
    mx = (0, 0, 0, 0)  # biggest bounding box so far
    mx_area = 0
    for cont in contours:
        x, y, w, h = cv2.boundingRect(cont)
        area = w*h
        if area > mx_area:
            mx = x, y, w, h
            mx_area = area
    x, y, w, h = mx

# Draw contours and crop image
    image_copy = img.copy()
    cv2.drawContours(image=image_copy, contours=contours, contourIdx=-1, color=(0, 255, 0), thickness=2, lineType=cv2.LINE_AA)
    roi = img[y:y+h, x:x+w]
    cv2.imwrite('Image_crop.jpg', roi)
    # cv2.imshow('Image', roi)
    # cv2.waitKey(0)

# Perform OCR on cropped image
    img_pil = Image.fromarray(roi)
    text = pytesseract.image_to_string(img_pil, lang='eng+tel')

# Convert text to hash
    hash_object = hashlib.sha256(text.encode())
    hex_dig = hash_object.hexdigest()
    return hex_dig
  

    # hashdoc = docx.Document()
    # hashdoc.add_paragraph(hex_dig)
    # hashdoc.save('OCR Hash.docx')

def generate_hashi(path):
    # Convert text to hash
    # img = np.array(img1)
    path1= './uploads-i/'+path



# Set Tesseract path
    os.environ['TESSDATA_PREFIX'] = 'C:/Users/Mahi Singhal/AppData/Local/Programs/Tesseract-OCR/tessdata'
    pytesseract.pytesseract.tesseract_cmd = r'C:/Users/Mahi Singhal/AppData/Local/Programs/Tesseract-OCR/tesseract.exe'

# Load image
    img = cv2.imread(path1)

# Convert image to grayscale
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

# Equalize histogram to enhance contrast
    equalized = cv2.equalizeHist(gray)

# Resize image
    new_size = (int(img.shape[1]*0.35), int(img.shape[0]*0.25))
    resized = cv2.resize(img, new_size) 

# Adjust brightness and contrast
    alpha = 1.25  # scaling factor for brightness adjustment
    beta = 30  # scalar added to the result
    brightened = cv2.convertScaleAbs(resized, alpha=alpha, beta=beta)

# Detect contours and find object with the biggest bounding box
    edges = cv2.Canny(gray, 50, 200)
    contours, hierarchy = cv2.findContours(edges.copy(), cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_NONE)
    mx = (0, 0, 0, 0)  # biggest bounding box so far
    mx_area = 0
    for cont in contours:
        x, y, w, h = cv2.boundingRect(cont)
        area = w*h
        if area > mx_area:
            mx = x, y, w, h
            mx_area = area
    x, y, w, h = mx

# Draw contours and crop image
    image_copy = img.copy()
    cv2.drawContours(image=image_copy, contours=contours, contourIdx=-1, color=(0, 255, 0), thickness=2, lineType=cv2.LINE_AA)
    roi = img[y:y+h, x:x+w]
    cv2.imwrite('Image_crop.jpg', roi)
    # cv2.imshow('Image', roi)
    # cv2.waitKey(0)

# Perform OCR on cropped image
    img_pil = Image.fromarray(roi)
    text = pytesseract.image_to_string(img_pil, lang='eng+tel')

# Convert text to hash
    hash_object = hashlib.sha256(text.encode())
    hex_dig = hash_object.hexdigest()
    return hex_dig
  

    # hashdoc = docx.Document()
    # hashdoc.add_paragraph(hex_dig)
    # hashdoc.save('OCR Hash.docx')

